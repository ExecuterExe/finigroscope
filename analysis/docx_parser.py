"""Чтение .docx: документ → список абзацев с разрешённым форматированием.

python-docx даёт иерархию Document → Paragraph → Run. Главная тонкость —
наследование: `run.bold = None` означает «наследует от стиля», а не «не жирный»
(см. раздел 4.1 проектного документа). Поэтому эффективные полужирность и шрифт
вычисляются с учётом стиля абзаца.
"""

import re
from dataclasses import dataclass
from typing import Optional

from docx import Document

from analysis.text_utils import collapse_spaces

# Доля «жирных» символов в абзаце, при которой считаем заголовок полужирным.
_BOLD_RATIO = 0.6


@dataclass
class Para:
    """Один значимый абзац документа с разрешённым форматированием."""

    index: int
    text: str               # очищенный текст (одинарные пробелы)
    style_name: str         # имя стиля абзаца
    is_heading_style: bool  # оформлен ли стилем «Заголовок …» / «Heading …»
    bold: bool              # эффективная полужирность (с учётом стиля)
    font_name: Optional[str]  # эффективный шрифт (первый значимый), None — неизвестен
    alignment: Optional[int]  # выравнивание: 0 left, 1 center, 2 right, 3 justify, None — наследует
    heading_level: Optional[int]  # уровень заголовка (1,2,…) из стиля, иначе None


def _style_bold(paragraph) -> Optional[bool]:
    try:
        return paragraph.style.font.bold
    except Exception:
        return None


def _style_font(paragraph) -> Optional[str]:
    try:
        return paragraph.style.font.name
    except Exception:
        return None


def _effective_bold(paragraph) -> bool:
    """Полужирность по доле символов: run.bold=None наследуется от стиля."""
    style_bold = _style_bold(paragraph)
    bold_chars = 0
    total_chars = 0
    for run in paragraph.runs:
        n = len(run.text.strip())
        if n == 0:
            continue
        total_chars += n
        resolved = style_bold if run.bold is None else run.bold
        if resolved:
            bold_chars += n
    if total_chars == 0:
        return bool(style_bold)
    return bold_chars / total_chars >= _BOLD_RATIO


def _effective_font(paragraph) -> Optional[str]:
    """Первый явно заданный шрифт среди значимых ранов, иначе шрифт стиля."""
    for run in paragraph.runs:
        if run.text.strip() and run.font.name:
            return run.font.name
    return _style_font(paragraph)


def _is_heading_style(style_name: str) -> bool:
    s = (style_name or "").lower()
    return s.startswith("heading") or s.startswith("заголовок")


def heading_level(style_name: str) -> Optional[int]:
    """Уровень заголовка из имени стиля: «Heading 2»/«Заголовок 2» → 2."""
    m = re.search(r"(?:heading|заголовок)\s*(\d+)", (style_name or "").lower())
    return int(m.group(1)) if m else None


def _alignment(paragraph) -> Optional[int]:
    try:
        a = paragraph.alignment
        return int(a) if a is not None else None
    except Exception:
        return None


def parse(source) -> list:
    """Читает .docx (путь или файловый объект) → список непустых абзацев Para."""
    document = Document(source)
    paras = []
    for i, paragraph in enumerate(document.paragraphs):
        text = collapse_spaces(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        paras.append(
            Para(
                index=i,
                text=text,
                style_name=style_name,
                is_heading_style=_is_heading_style(style_name),
                bold=_effective_bold(paragraph),
                font_name=_effective_font(paragraph),
                alignment=_alignment(paragraph),
                heading_level=heading_level(style_name),
            )
        )
    return paras


_EMU_PER_CM = 360000


def document_meta(source) -> dict:
    """Документные свойства для строгой проверки эссе: поля страницы (см)."""
    meta = {"margins": None}
    try:
        document = Document(source)
        sec = document.sections[0]

        def cm(v):
            return round(v / _EMU_PER_CM, 2) if v is not None else None

        meta["margins"] = {
            "top": cm(sec.top_margin), "bottom": cm(sec.bottom_margin),
            "left": cm(sec.left_margin), "right": cm(sec.right_margin),
        }
    except Exception:
        pass
    return meta
