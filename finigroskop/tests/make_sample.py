"""Генератор синтетического .docx для проверки этапа 1.

Собирает документ с НЕСКОЛЬКИМИ играми (повтор структуры эссе) с намеренными
дефектами: пропущенный раздел, не-полужирный заголовок, чужой шрифт, короткий
раздел — чтобы проверить, что движок их ловит. Возвращает BytesIO.
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt

from analysis.reference import STRUCTURE_ESSAY

LOREM = ("Это содержательный абзац раздела с достаточным объёмом текста, чтобы "
         "пройти проверку наполненности. Здесь описывается соответствующая часть "
         "игры подробно и по делу, без воды и повторов одного и того же.")


def _add_heading(doc, text, bold=True, font="Times New Roman", heading_style=True):
    style = "Heading 2" if (heading_style and "." in text.split()[0]) else (
        "Heading 1" if heading_style else "Normal")
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(13)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def build_sample() -> BytesIO:
    doc = Document()

    # --- Игра 1: почти полная, с извлекаемыми параметрами ---
    for i, head in enumerate(STRUCTURE_ESSAY):
        if head.startswith("4.6"):
            continue  # намеренно пропущенный раздел
        bold = not head.startswith("2.2")          # 2.2 — не полужирный (дефект)
        font = "Calibri" if head.startswith("3.1") else "Times New Roman"  # 3.1 — чужой шрифт
        _add_heading(doc, head, bold=bold, font=font)
        # родительские разделы (1., 2., ...) оставляем без текста — это норма
        if "." in head.split()[0]:
            body = LOREM
            if head.startswith("1.4"):
                body = "Игра рассчитана на 2-4 игрока, возраст 12+. " + LOREM
            if head.startswith("4.1"):
                body = ("Побеждает тот, кто первым доберётся до финиша. Игра "
                        "заканчивается, когда первый игрок пересекает финишную черту. " + LOREM)
            if head.startswith("3.1"):
                body = "В колоде 50 карточек вопросов и колода из 30 карт событий. " + LOREM
            if head.startswith("4.3"):
                body = "Игрок бросает кубик d6 и двигается. За верный ответ +4 балла, за неверный -2 очка. " + LOREM
            if head.startswith("5.2"):
                body = "Партия рассчитана на 20-40 минут. " + LOREM
            _add_body(doc, body)

    # --- Игра 2: заметно слабее, без параметров, с короткими разделами ---
    for head in STRUCTURE_ESSAY:
        if head.startswith(("4.", "5.", "6.")) and "." in head.split()[0][1:]:
            continue  # у второй игры выпала половина подразделов
        _add_heading(doc, head, bold=True, font="Times New Roman")
        if "." in head.split()[0]:
            _add_body(doc, "Кратко.")  # слишком мало текста

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


if __name__ == "__main__":
    out = build_sample()
    with open("tests/sample.docx", "wb") as f:
        f.write(out.read())
    print("Записан tests/sample.docx")
